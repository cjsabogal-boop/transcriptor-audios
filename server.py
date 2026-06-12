import os
import sys
import threading
import multiprocessing as mp
import subprocess
import time
import json
import requests

# Intentar cargar .env si existe para la clave de API de Gemini
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configuración GLOBAL de estabilidad para Mac (antes de importar torch)

# Configuración GLOBAL de estabilidad para Mac (antes de importar torch)
if sys.platform == "darwin":
    os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
    os.environ["PYTORCH_MPS_HIGH_WATERMARK_RATIO"] = "0.0"

try:
    import torch
    import whisper
    if torch is not None:
        torch.set_num_threads(4)
        if sys.platform == "darwin" and hasattr(torch.backends, "mps"):
            torch.backends.mps.is_available = lambda: False
except ImportError:
    torch = None
    whisper = None

# Motor preferido: faster-whisper (CTranslate2). Más rápido y mucho menos RAM
# (int8 en CPU). Si no está instalado, se cae a openai-whisper.
try:
    from faster_whisper import WhisperModel as FasterWhisperModel
    FW_AVAILABLE = True
except ImportError:
    FasterWhisperModel = None
    FW_AVAILABLE = False

from flask import Flask, jsonify, request, send_file, send_from_directory
try:
    from flask_cors import CORS
except ImportError:
    CORS = None
from werkzeug.utils import secure_filename

app = Flask(__name__, static_folder="frontend/public", static_url_path="")

@app.route("/")
def index():
    """Servimos el frontend desde la misma aplicación Flask."""
    return send_from_directory(app.static_folder, "index.html")

# Referencias globales para el estado y el proceso
STATE = None
manager = None
CURRENT_PROCESS = None

# FALLBACK DE RUTAS PARA MAC
_LOCAL_BIN = os.path.expanduser("~/.local/bin")
_EXTRA_PATHS = [_LOCAL_BIN, "/opt/homebrew/bin", "/usr/local/bin"]
for _p in reversed(_EXTRA_PATHS):
    if _p not in os.environ.get("PATH", ""):
        os.environ["PATH"] = _p + os.pathsep + os.environ.get("PATH", "")

def find_binary(name):
    import shutil
    local_path = os.path.join(_LOCAL_BIN, name)
    if os.path.isfile(local_path) and os.access(local_path, os.X_OK):
        return local_path
    binary_path = shutil.which(name)
    if binary_path:
        return binary_path
    standard_paths = ["/opt/homebrew/bin", "/usr/local/bin", "/usr/bin"]
    for p in standard_paths:
        full_path = os.path.join(p, name)
        if os.path.isfile(full_path) and os.access(full_path, os.X_OK):
            return full_path
    return name

@app.after_request
def no_cache(resp):
    # App local: nunca cachear, así el navegador siempre carga la última versión
    # (evita ver pantallas viejas tras actualizar el frontend).
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


if CORS:
    CORS(app, resources={r"/api/*": {"origins": "*"}})
else:
    @app.after_request
    def add_cors_headers(response):
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'GET, POST, DELETE, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response

AUDIO_DIR = os.path.expanduser("~/Documents/Audios")
os.makedirs(AUDIO_DIR, exist_ok=True)
app.config['UPLOAD_FOLDER'] = AUDIO_DIR
SERVER_START_TIME = time.time()

# ── Apagado automático al cerrar la ventana (libera la RAM/modelo) ──
# Pensado para Macs con poca RAM: si no hay ventana abierta, el servidor se apaga.
AUTO_SHUTDOWN = os.environ.get("AUTO_SHUTDOWN", "1") != "0"
CLOSE_GRACE = 6        # seg tras avisar de cierre (sin reabrir) -> apagar
BACKSTOP_GRACE = 120   # seg sin señales de la ventana -> apagar (respaldo)
LAST_SEEN = time.time()
SHUTDOWN_AT = None


def _watchdog():
    """Apaga el servidor cuando la ventana se cierra (nunca durante una transcripción)."""
    global SHUTDOWN_AT
    while True:
        time.sleep(2)
        try:
            if STATE and STATE.get("is_processing"):
                continue  # jamás apagar a mitad de un audio
            now = time.time()
            if (SHUTDOWN_AT and now >= SHUTDOWN_AT) or (now - LAST_SEEN > BACKSTOP_GRACE):
                print("🛑 Ventana cerrada: apagando para liberar memoria.")
                os._exit(0)
        except Exception:
            pass

# ── Configuración de modelos Whisper ──
# De más rápido/ligero a más preciso/pesado.
ALLOWED_MODELS = ["tiny", "base", "small", "medium", "large-v3"]
MODEL_INFO = {
    "tiny":     {"label": "Tiny — rápido, calidad básica",       "ram_gb": 1},
    "base":     {"label": "Base — equilibrado ligero",           "ram_gb": 1},
    "small":    {"label": "Small — buena calidad",               "ram_gb": 2},
    "medium":   {"label": "Medium — muy buena (recomendado)",    "ram_gb": 5},
    "large-v3": {"label": "Large-v3 — máxima calidad",           "ram_gb": 10},
}
# Modelo por defecto: lo define la variable de entorno WHISPER_MODEL.
# Con faster-whisper, "small" corre cómodo incluso en 8GB, así que es el baseline.
# El paquete para la máquina grande exporta WHISPER_MODEL=medium (lo sobreescribe).
DEFAULT_MODEL = os.environ.get("WHISPER_MODEL", "small").strip().lower()
if DEFAULT_MODEL not in ALLOWED_MODELS:
    DEFAULT_MODEL = "small"

# ── Auto-actualización desde GitHub ──
# Usamos la API de contenidos (con Accept raw) en vez de raw.githubusercontent.com
# para evitar el caché de CDN y que las actualizaciones sean inmediatas.
GITHUB_API = "https://api.github.com/repos/cjsabogal-boop/transcriptor-audios/contents"
UPDATE_FILES = [
    "server.py",
    "run_server.sh",
    "requirements.txt",
    "frontend/public/index.html",
    "frontend/public/app.js",
    "frontend/public/styles.css",
    "frontend/public/404.html",
]


def detectar_device():
    if torch is not None and getattr(torch, "cuda", None) is not None and torch.cuda.is_available():
        return "cuda"
    return "cpu"


def detectar_ram_gb():
    try:
        return round(os.sysconf("SC_PAGE_SIZE") * os.sysconf("SC_PHYS_PAGES") / (1024 ** 3), 1)
    except Exception:
        return None

# ── Singleton del modelo Whisper ──
WHISPER_MODELS = {}
WHISPER_DEVICE = "cpu"

def cargar_modelo(model_size="base"):
    """Devuelve una tupla (engine, modelo) donde engine es 'fw' (faster-whisper)
    u 'ow' (openai-whisper), para saber qué API usar al transcribir."""
    global WHISPER_MODELS, WHISPER_DEVICE
    if model_size in WHISPER_MODELS:
        return WHISPER_MODELS[model_size]

    device = "cuda" if (torch is not None and getattr(torch, "cuda", None) is not None and torch.cuda.is_available()) else "cpu"
    WHISPER_DEVICE = device

    if FW_AVAILABLE:
        # int8 en CPU (rápido y poca RAM); float16 en GPU.
        compute_type = "float16" if device == "cuda" else "int8"
        print(f"📦 Cargando faster-whisper '{model_size}' en {device.upper()} ({compute_type})...")
        modelo = FasterWhisperModel(model_size, device=device, compute_type=compute_type)
        WHISPER_MODELS[model_size] = ("fw", modelo)
        print(f"✅ Modelo '{model_size}' cargado (faster-whisper).")
    else:
        if sys.platform == "darwin" and torch is not None:
            torch.set_num_threads(mp.cpu_count())
        print(f"📦 Cargando openai-whisper '{model_size}' en {device.upper()}...")
        WHISPER_MODELS[model_size] = ("ow", whisper.load_model(model_size, device=device))
        print(f"✅ Modelo '{model_size}' cargado (openai-whisper).")

    return WHISPER_MODELS[model_size]

def obtener_duracion_audio(ruta):
    try:
        bin_ffprobe = find_binary('ffprobe')
        result = subprocess.run(
            [bin_ffprobe, '-v', 'quiet', '-show_entries', 'format=duration',
             '-of', 'default=noprint_wrappers=1:nokey=1', ruta],
            capture_output=True, text=True, timeout=10
        )
        return float(result.stdout.strip())
    except Exception:
        return None

def preconvertir_audio(ruta_entrada, start_time=None, end_time=None):
    """Convierte audio al formato que Whisper necesita (mono, 16kHz). Rápido, sin filtros pesados."""
    suffix = "_optimized.wav"
    if start_time is not None or end_time is not None:
        suffix = f"_trim_{int(start_time or 0)}_{int(end_time or 0)}.wav"
    
    ruta_opt = ruta_entrada.rsplit('.', 1)[0] + suffix
    try:
        bin_ffmpeg = find_binary('ffmpeg')
        cmd = [bin_ffmpeg, '-y', '-nostdin']
        
        # Recorte si se solicita
        if start_time: cmd.extend(['-ss', str(start_time)])
        if end_time: cmd.extend(['-to', str(end_time)])
        
        cmd.extend([
            '-i', ruta_entrada,
            '-ac', '1', '-ar', '16000',
            '-c:a', 'pcm_s16le',
            '-loglevel', 'error', ruta_opt
        ])
        
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE,
                       stdin=subprocess.DEVNULL, timeout=120)
        
        if os.path.exists(ruta_opt) and os.path.getsize(ruta_opt) > 0:
            return ruta_opt
        return ruta_entrada
    except Exception as e:
        print(f"⚠️ Pre-conversión falló: {e}")
        return ruta_entrada

def limpiar_temporales():
    try:
        for f in os.listdir(AUDIO_DIR):
            if f.endswith('.wav') and ('_optimized' in f or '_trim_' in f or '_temp' in f):
                os.remove(os.path.join(AUDIO_DIR, f))
    except Exception:
        pass

def check_dependencies():
    # Falta instalar si no hay NINGÚN motor de transcripción disponible.
    return not (FW_AVAILABLE or whisper is not None)


@app.route("/api/health")
def health():
    # Cada health-check de la ventana cuenta como "latido": mantiene vivo el server
    # y cancela cualquier apagado pendiente (p. ej. si fue solo un recargar la página).
    global LAST_SEEN, SHUTDOWN_AT
    LAST_SEEN = time.time()
    SHUTDOWN_AT = None
    return jsonify({
        "status": "ok",
        "uptime": round(time.time() - SERVER_START_TIME),
        "needs_install": check_dependencies(),
        "is_processing": STATE.get("is_processing", False) if STATE else False
    })


@app.route("/api/closing", methods=["POST", "GET"])
def api_closing():
    """La ventana avisa que se está cerrando (navigator.sendBeacon). Agenda el
    apagado; si fue un recargar, el próximo /api/health lo cancela."""
    global SHUTDOWN_AT
    SHUTDOWN_AT = time.time() + CLOSE_GRACE
    return ("", 204)

@app.route("/api/config")
def api_config():
    """Info de modelos disponibles y de la máquina, para que el frontend
    arme el selector de modelo y muestre el hardware."""
    return jsonify({
        "default_model": DEFAULT_MODEL,
        "allowed_models": ALLOWED_MODELS,
        "model_info": MODEL_INFO,
        "loaded_models": list(WHISPER_MODELS.keys()),
        "engine": "faster-whisper" if FW_AVAILABLE else "openai-whisper",
        "device": detectar_device(),
        "total_ram_gb": detectar_ram_gb(),
        "platform": sys.platform,
    })


@app.route("/api/status")
def status():
    global CURRENT_PROCESS
    res = dict(STATE or {})
    if res.get("is_processing") and CURRENT_PROCESS and not CURRENT_PROCESS.is_alive() and not res.get("completed"):
        res["is_processing"] = False
        res["error"] = True
        res["status_text"] = "❌ El proceso de IA se detuvo inesperadamente."
        STATE.update(res)
    return jsonify(res)

@app.route("/api/audios")
def get_audios():
    audios = []
    try:
        text_files = [f for f in os.listdir(AUDIO_DIR) if f.lower().endswith('.txt') and f != "logs_sistema.txt"]
        for txt in text_files:
            bname = os.path.splitext(txt)[0]
            txt_path = os.path.join(AUDIO_DIR, txt)
            json_path = os.path.join(AUDIO_DIR, bname + ".json")
            fecha_mod = os.path.getmtime(txt_path)
            
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as fj:
                    data = json.load(fj)
                    text_es = data.get("text_es", "")
                    text_en = data.get("text_en", "")
                    size_bytes = len(text_es.encode('utf-8')) + len(text_en.encode('utf-8'))
            else:
                with open(txt_path, 'r', encoding='utf-8') as f:
                    text_es = f.read()
                    text_en = ""
                    size_bytes = len(text_es.encode('utf-8'))
            
            audios.append({
                "title": bname, 
                "text_es": text_es, 
                "text_en": text_en, 
                "date": fecha_mod, 
                "size": size_bytes
            })
        audios.sort(key=lambda x: x['date'], reverse=True)
    except Exception: pass
    return jsonify(audios)

@app.route("/api/upload", methods=["POST"])
def upload_file():
    global CURRENT_PROCESS
    if STATE.get("is_processing"):
        return jsonify({"success": False, "msg": "Ya hay una transcripción en proceso."})
    
    if 'audiofile' not in request.files:
        return jsonify({"success": False, "msg": "No se envió audio"})
    
    file = request.files['audiofile']
    start_time = request.form.get('start_time')
    end_time = request.form.get('end_time')
    language = request.form.get('language', 'es')  # Default español
    task = request.form.get('task', 'transcribe')   # 'transcribe' | 'translate'
    if task not in ("transcribe", "translate"):
        task = "transcribe"
    # Modelo solicitado (override por subida) o el default del servidor
    model_size = (request.form.get('model') or DEFAULT_MODEL).strip().lower()
    if model_size not in ALLOWED_MODELS:
        model_size = DEFAULT_MODEL

    # Convertir a float si existen
    try:
        start_time = float(start_time) if start_time else None
        end_time = float(end_time) if end_time else None
    except ValueError:
        start_time = end_time = None

    filename = secure_filename(file.filename)
    if start_time is not None or end_time is not None:
        filename = f"part_{int(start_time or 0)}_{int(end_time or 0)}_{filename}"

    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(save_path)

    if CURRENT_PROCESS and CURRENT_PROCESS.is_alive(): CURRENT_PROCESS.terminate()
    CURRENT_PROCESS = mp.Process(target=proceso_fondo, args=(save_path, filename, STATE, start_time, end_time, language, model_size, task))
    CURRENT_PROCESS.start()
    return jsonify({"success": True})

def proceso_fondo(ruta_audio, filename, estado, start_time, end_time, language, model_size, task):
    import traceback
    def log(msg):
        with open("child_debug.log", "a") as f: f.write(f"[{time.ctime()}] {msg}\n")
    
    try:
        estado["is_processing"] = True
        estado["completed"] = False
        estado["error"] = False
        estado["progress_percent"] = 5.0
        estado["current_filename"] = filename
        estado["status_text"] = "Convirtiendo audio..."
        
        duracion = obtener_duracion_audio(ruta_audio)
        base = os.path.splitext(filename)[0]
        
        # Medir tamaño ANTES de convertir (el archivo original comprimido)
        size_mb_original = os.path.getsize(ruta_audio) / (1024 * 1024)
        
        estado["progress_percent"] = 10.0
        estado["status_text"] = "Preparando audio para IA..."
        
        # 1. Pre-procesar (conversión rápida, sin filtros pesados)
        ruta_procesable = preconvertir_audio(ruta_audio, start_time, end_time)
        ruta_wav_temp = ruta_procesable if ruta_procesable != ruta_audio else None
        
        # 2. Transcribir con el modelo elegido (configurable por el usuario / servidor)
        estado["progress_percent"] = 20.0
        estado["status_text"] = f"Cargando modelo '{model_size}'..."
        engine, modelo = cargar_modelo(model_size)

        # LOG para depuración
        print(f"📊 Depuración - Idioma: '{language}', Tarea: '{task}', Archivo: {size_mb_original:.1f}MB, Modelo: {model_size}, Motor: {engine}")

        # Idioma de origen para Whisper. Vacío/None => auto-detectar.
        if not language or language.strip() == "" or language == "None":
            lang_whisper = None
        else:
            lang_whisper = language.strip().lower()
            if lang_whisper not in ["es", "en"]:
                lang_whisper = None

        verbo = "Traduciendo" if task == "translate" else "Transcribiendo"
        estado["status_text"] = f"{verbo} ({lang_whisper or 'Auto'}) · modelo {model_size}..."
        estado["progress_percent"] = 30.0

        # Transcripción directa (Whisper siempre transcribe en el idioma original;
        # la traducción a español, si se pide, la hace Gemini después en el frontend).
        if engine == "fw":
            seg_iter, info = modelo.transcribe(ruta_procesable, language=lang_whisper, beam_size=5)
            partes, segments = [], []
            for i, s in enumerate(seg_iter):
                partes.append(s.text)
                segments.append({"id": i, "start": s.start, "end": s.end, "text": s.text})
            texto = "".join(partes).strip()
        else:
            resultado = modelo.transcribe(ruta_procesable, language=lang_whisper, fp16=False)
            texto = resultado["text"].strip()
            segments = resultado.get("segments", [])

        # 3. Guardar resultados.
        # text_en = transcripción cruda original (entrada para la IA).
        # text_es = salida editable: si es transcripción la mostramos tal cual;
        #           si es traducción la deja vacía para que Gemini la rellene.
        text_es_val = "" if task == "translate" else texto
        with open(os.path.join(AUDIO_DIR, base + ".json"), 'w', encoding='utf-8') as f:
            json.dump({
                "text_es": text_es_val,
                "text_en": texto,
                "segments": segments,
                "language": language,
                "task": task,
                "model": model_size,
                "trimmed": (start_time is not None or end_time is not None)
            }, f, ensure_ascii=False, indent=2)
            
        with open(os.path.join(AUDIO_DIR, base + ".txt"), 'w', encoding='utf-8') as f:
            f.write(texto)

        estado["progress_percent"] = 100.0
        estado["completed"] = True
        estado["title"] = base  # Para auto-abrir el modal en el frontend
        estado["task"] = task    # 'translate' => el frontend lanza Gemini automáticamente
        estado["status_text"] = "✅ ¡Listo!"
                    
    except Exception as e:
        log(f"ERROR: {str(e)}\n{traceback.format_exc()}")
        estado["status_text"] = f"❌ Error: {str(e)}"
        estado["error"] = True
    finally:
        if 'ruta_wav_temp' in locals() and ruta_wav_temp and os.path.exists(ruta_wav_temp):
            try: os.remove(ruta_wav_temp)
            except: pass
        estado["is_processing"] = False

@app.route("/api/audios/<title>/save", methods=["POST"])
def save_audio(title):
    data = request.json
    text_es = data.get("text_es", "")
    safe_title = secure_filename(title)
    json_path = os.path.join(AUDIO_DIR, safe_title + ".json")
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f: jdata = json.load(f)
        jdata["text_es"] = text_es
        with open(json_path, 'w', encoding='utf-8') as f: json.dump(jdata, f, ensure_ascii=False, indent=2)
        with open(os.path.join(AUDIO_DIR, safe_title + ".txt"), 'w', encoding='utf-8') as f: f.write(text_es)
        return jsonify({"success": True})
    return jsonify({"success": False}), 404


@app.route("/api/audios/<title>/segments")
def get_segments(title):
    safe_title = secure_filename(title)
    json_path = os.path.join(AUDIO_DIR, safe_title + ".json")
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return jsonify(data.get("segments", []))
    return jsonify([]), 404

@app.route("/api/audios/<title>", methods=["DELETE"])
def delete_audio(title):
    safe_title = secure_filename(title)
    # Lista de archivos a borrar
    files_to_delete = [
        safe_title + '.json',
        safe_title + '.txt',
        safe_title + '.mp3',
        safe_title + '.wav',
        safe_title + '.m4a',
        safe_title + '_audio_es.mp3'
    ]
    for filename in files_to_delete:
        filepath = os.path.join(AUDIO_DIR, filename)
        if os.path.exists(filepath): os.remove(filepath)
    return jsonify({"success": True})

@app.route("/api/audios/<title>/rename", methods=["POST"])
def rename_audio(title):
    data = request.json
    new_title = data.get("new_title", "")
    if not new_title: return jsonify({"success": False, "msg": "Nuevo nombre requerido"}), 400
    
    old_safe = secure_filename(title)
    new_safe = secure_filename(new_title)
    
    if os.path.exists(os.path.join(AUDIO_DIR, new_safe + ".json")):
        return jsonify({"success": False, "msg": "Ya existe un archivo con ese nombre"}), 400

    # Extensiones a renombrar
    exts = ['.json', '.txt', '.mp3', '.wav', '.m4a', '_audio_es.mp3']
    for ext in exts:
        old_path = os.path.join(AUDIO_DIR, (old_safe + ext) if ext.startswith('.') else (old_safe + ext))
        new_path = os.path.join(AUDIO_DIR, (new_safe + ext) if ext.startswith('.') else (new_safe + ext))
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
            
    # También actualizar el título dentro del JSON
    json_path = os.path.join(AUDIO_DIR, new_safe + ".json")
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            jdata = json.load(f)
        jdata["title"] = new_title
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(jdata, f, ensure_ascii=False, indent=2)

    return jsonify({"success": True})
@app.route("/api/install", methods=["POST"])
def install_dependencies():
    """Lanza la instalación de dependencias en un hilo separado."""
    def run_install():
        STATE["is_installing"] = True
        STATE["status_text"] = "Iniciando instalación de IA..."
        try:
            # Whisper + PyTorch + EdgeTTS + Google Generative AI
            deps = ["openai-whisper", "torch", "torchvision", "torchaudio", "edge-tts", "google-generativeai"]
            total = len(deps)
            for i, dep in enumerate(deps):
                STATE["status_text"] = f"Instalando {dep} ({i+1}/{total})..."
                subprocess.run([sys.executable, "-m", "pip", "install", dep], check=True)
            STATE["status_text"] = "✅ Instalación finalizada."
        except Exception as e:
            STATE["status_text"] = f"❌ Error en instalación: {str(e)}"
        finally:
            STATE["is_installing"] = False

    t = threading.Thread(target=run_install)
    t.start()
    return jsonify({"success": True})

PROMPTS_POR_MODO = {
    "general": (
        "Eres un editor profesional. Toma la siguiente transcripción cruda de un audio "
        "(puede ser un podcast, charla, conferencia, entrevista, audio personal, etc.) y "
        "produce una versión limpia, fluida y profesional en ESPAÑOL.\n\n"
        "INSTRUCCIONES:\n"
        "1. ELIMINA muletillas, repeticiones, palabras de relleno, falsos comienzos y silencios verbales (eh, ah, este, o sea, pues, ¿no?).\n"
        "2. CONSERVA el contenido sustantivo: ideas, datos, ejemplos, anécdotas, conclusiones.\n"
        "3. Si el original está en inglés u otro idioma, TRADÚCELO al español natural.\n"
        "4. Mejora la fluidez para que se lea de corrido, manteniendo el tono del autor.\n"
        "5. Conserva los nombres propios, cifras y citas tal cual.\n"
        "6. Si hay claramente varios bloques temáticos, separa con párrafos.\n"
    ),
    "interview": (
        "Eres un editor de entrevistas. Toma la siguiente transcripción cruda y produce "
        "una versión limpia en formato de diálogo en ESPAÑOL.\n\n"
        "INSTRUCCIONES:\n"
        "1. Detecta cambios de hablante e identifícalos como 'Entrevistador:' y 'Invitado:' (o nombres si los menciona).\n"
        "2. Elimina muletillas, repeticiones y falsos comienzos.\n"
        "3. Traduce al español si el original está en otro idioma.\n"
        "4. Mantén las preguntas y respuestas completas, sin perder ideas importantes.\n"
        "5. Cada intervención en un párrafo separado.\n"
    ),
    "summary": (
        "Eres un editor experto en sintetizar. Toma la siguiente transcripción y produce "
        "un RESUMEN EJECUTIVO en ESPAÑOL.\n\n"
        "INSTRUCCIONES:\n"
        "1. Empieza con 1-2 oraciones que capturen la idea principal.\n"
        "2. Sigue con 3-7 puntos clave en párrafos cortos (no uses viñetas ni markdown).\n"
        "3. Cierra con una conclusión o llamado a la acción si lo hay.\n"
        "4. Traduce al español si el original está en otro idioma.\n"
        "5. Mantén nombres propios, cifras y citas relevantes.\n"
    ),
    "sermon": (
        "Eres un editor profesional de contenido cristiano en español. Toma la siguiente "
        "transcripción de un sermón y produce un guion de audio fluido, claro y profesional en ESPAÑOL.\n\n"
        "INSTRUCCIONES OBLIGATORIAS:\n"
        "1. ELIMINA completamente: anuncios iniciales, letras de canciones, avisos finales, saludos de bienvenida al servicio, instrucciones logísticas.\n"
        "2. ENFÓCATE EXCLUSIVAMENTE en el mensaje bíblico/sermón principal.\n"
        "3. Si el texto original está en inglés, tradúcelo al español de forma natural y fluida.\n"
        "4. Adapta el texto para que suene natural al ser LEÍDO EN VOZ ALTA (es un guion de audio).\n"
        "5. Mantén las citas bíblicas y referencias escriturales.\n"
        "6. Usa un tono cálido, cercano y pastoral.\n"
    ),
    "raw": (
        "Eres un editor literal. Toma la siguiente transcripción cruda y produce una versión "
        "que conserve EXACTAMENTE el contenido, solo con correcciones mínimas en ESPAÑOL.\n\n"
        "INSTRUCCIONES:\n"
        "1. Corrige solo errores evidentes de transcripción (palabras mal entendidas, puntuación).\n"
        "2. NO elimines muletillas, NO reformules, NO resumas.\n"
        "3. Si el original está en otro idioma, traduce literalmente.\n"
        "4. Mantén el orden y todas las palabras del original.\n"
    ),
}

REGLAS_FORMATO_COMUNES = (
    "\nREGLAS DE FORMATO:\n"
    "- Devuelve ÚNICAMENTE el texto adaptado, listo para leer.\n"
    "- NO incluyas introducciones como 'Aquí tienes...', 'Claro...', 'A continuación...'.\n"
    "- NO uses formato Markdown (**, ##, ---, viñetas).\n"
    "- Solo texto puro, párrafos separados por líneas en blanco.\n\n"
)

def construir_prompt(mode, texto_original):
    base = PROMPTS_POR_MODO.get(mode, PROMPTS_POR_MODO["general"])
    return base + REGLAS_FORMATO_COMUNES + f"TRANSCRIPCIÓN ORIGINAL:\n{texto_original}"


@app.route("/api/gemini/transform", methods=["POST"])
def gemini_transform():
    import traceback
    try:
        import google.generativeai as genai
        data = request.json
        text_en = data.get("text_en", "")
        title = data.get("title", "Audio")
        mode = data.get("mode", "general")
        if mode not in PROMPTS_POR_MODO:
            mode = "general"

        if not text_en:
            # Intentar cargar desde el archivo si no viene en el body
            safe_title = secure_filename(title)
            json_path = os.path.join(AUDIO_DIR, safe_title + ".json")
            if os.path.exists(json_path):
                try:
                    with open(json_path, 'r', encoding='utf-8') as f:
                        jdata = json.load(f)
                        text_en = jdata.get("text_en", "")
                except Exception as e:
                    print(f"⚠️ Error al leer JSON para Gemini: {e}")

        if not text_en:
            return jsonify({"success": False, "msg": "No hay texto para procesar."}), 400

        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            return jsonify({
                "success": False,
                "msg": "Falta la clave de API de Gemini (GOOGLE_API_KEY). Configúrala en el servidor."
            }), 401

        print(f"✨ Adaptación Gemini [{mode}] para: {title}")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.5-flash-lite')

        prompt = construir_prompt(mode, text_en)

        response = model.generate_content(prompt)
        text_es = response.text.replace('```markdown', '').replace('```', '').strip()

        print(f"✅ Adaptación exitosa [{mode}]: {title}")
        return jsonify({"success": True, "text_es": text_es, "mode": mode})
    except Exception as e:
        err_msg = traceback.format_exc()
        print(f"❌ Error CRÍTICO en Gemini:\n{err_msg}")
        return jsonify({"success": False, "msg": str(e)}), 500

# TTS endpoints se mantienen igual (omitidos por brevedad pero incluidos en la escritura final)
@app.route("/api/audios/<title>/tts", methods=["POST"])
def generate_tts(title):
    text_es = request.json.get("text_es", "")
    if not text_es: return jsonify({"success": False}), 400
    safe_title = secure_filename(title)
    output_mp3 = os.path.join(AUDIO_DIR, safe_title + "_audio_es.mp3")
    import edge_tts, asyncio
    
    async def _crear():
        # Gerardo es la voz colombiana masculina natural
        communicate = edge_tts.Communicate(text_es, "es-CO-GonzaloNeural", rate="-5%")
        await communicate.save(output_mp3)

    try:
        # Modo más robusto de correr asyncio en hilos de Flask
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(_crear())
        finally:
            loop.close()
            
        return jsonify({"success": True, "download_url": f"/api/download_tts/{safe_title}_audio_es.mp3"})
    except Exception as e:
        print(f"❌ Error TTS: {e}")
        return jsonify({"success": False, "msg": str(e)}), 500

@app.route("/api/download_tts/<filename>")
def download_tts_file(filename):
    from flask import send_from_directory
    return send_from_directory(AUDIO_DIR, filename, as_attachment=True)


@app.route("/api/translate", methods=["POST"])
def api_translate():
    """Traduce texto a español (u otro idioma) con Google Translate libre.
    No usa Gemini ni ninguna API key."""
    data = request.json or {}
    text = (data.get("text") or "").strip()
    target = (data.get("target") or "es").strip().lower()
    if not text:
        return jsonify({"success": False, "msg": "No hay texto para traducir."}), 400
    try:
        from deep_translator import GoogleTranslator
        # Google Translate limita ~5000 chars por llamada: troceamos por seguridad.
        chunks = [text[i:i + 4500] for i in range(0, len(text), 4500)]
        traducidos = [GoogleTranslator(source="auto", target=target).translate(c) for c in chunks]
        return jsonify({"success": True, "text": "\n".join(t for t in traducidos if t)})
    except Exception as e:
        print(f"❌ Error traduciendo: {e}")
        return jsonify({"success": False, "msg": str(e)}), 500


# ════════════════════════════════════════════════════════════════
# HERRAMIENTAS: subtítulos, exportar, buscar, URL, Ollama
# ════════════════════════════════════════════════════════════════

def _auto_pip(*paquetes):
    """Instala paquetes pip al vuelo si faltan (para usuarios que ya tenían
    la app instalada antes de esta versión)."""
    try:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", *paquetes, "--break-system-packages"],
            capture_output=True, timeout=300
        )
    except Exception:
        try:
            subprocess.run([sys.executable, "-m", "pip", "install", *paquetes],
                           capture_output=True, timeout=300)
        except Exception:
            pass


def _leer_json_audio(title):
    safe = secure_filename(title)
    p = os.path.join(AUDIO_DIR, safe + ".json")
    if not os.path.exists(p):
        return None, safe
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f), safe


def _fmt_ts(seg, vtt=False):
    """Segundos -> '00:01:23,456' (srt) o '00:01:23.456' (vtt)."""
    ms = int(round((seg - int(seg)) * 1000))
    s = int(seg)
    h, s = divmod(s, 3600)
    m, s = divmod(s, 60)
    sep = "." if vtt else ","
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


@app.route("/api/audios/<title>/subtitles")
def descargar_subtitulos(title):
    """Descarga subtítulos .srt o .vtt construidos con los timestamps de Whisper."""
    fmt = (request.args.get("fmt") or "srt").lower()
    data, safe = _leer_json_audio(title)
    if not data:
        return jsonify({"success": False, "msg": "Audio no encontrado."}), 404
    segments = data.get("segments") or []
    if not segments:
        return jsonify({"success": False, "msg": "Este audio no tiene marcas de tiempo (fue transcrito con una versión vieja). Vuelve a transcribirlo."}), 400

    lineas = []
    if fmt == "vtt":
        lineas.append("WEBVTT\n")
        for s in segments:
            lineas.append(f"{_fmt_ts(s['start'], vtt=True)} --> {_fmt_ts(s['end'], vtt=True)}\n{s['text'].strip()}\n")
        mime, ext = "text/vtt", "vtt"
    else:
        for i, s in enumerate(segments, 1):
            lineas.append(f"{i}\n{_fmt_ts(s['start'])} --> {_fmt_ts(s['end'])}\n{s['text'].strip()}\n")
        mime, ext = "application/x-subrip", "srt"

    import io
    buf = io.BytesIO("\n".join(lineas).encode("utf-8"))
    return send_file(buf, mimetype=mime, as_attachment=True, download_name=f"{safe}.{ext}")


@app.route("/api/audios/<title>/export")
def exportar_audio(title):
    """Exporta la transcripción a txt, docx (Word) o pdf."""
    fmt = (request.args.get("fmt") or "txt").lower()
    data, safe = _leer_json_audio(title)
    if not data:
        return jsonify({"success": False, "msg": "Audio no encontrado."}), 404
    texto = (data.get("text_es") or data.get("text_en") or "").strip()
    if not texto:
        return jsonify({"success": False, "msg": "No hay texto que exportar."}), 400
    fecha = time.strftime("%d/%m/%Y")
    # Título legible: los nombres de archivo largos con _ no se pueden partir en líneas
    titulo_legible = title.replace("_", " ").strip()
    import io

    if fmt == "txt":
        cuerpo = f"{titulo_legible}\n{'=' * len(titulo_legible)}\nFecha: {fecha}\n\n{texto}\n"
        buf = io.BytesIO(cuerpo.encode("utf-8"))
        return send_file(buf, mimetype="text/plain", as_attachment=True, download_name=f"{safe}.txt")

    if fmt == "docx":
        try:
            import docx  # python-docx
        except ImportError:
            _auto_pip("python-docx")
            try:
                import docx
            except ImportError:
                return jsonify({"success": False, "msg": "No se pudo instalar python-docx."}), 500
        doc = docx.Document()
        doc.add_heading(titulo_legible, level=1)
        doc.add_paragraph(f"Fecha: {fecha}").italic = True
        for par in texto.split("\n\n"):
            if par.strip():
                doc.add_paragraph(par.strip())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name=f"{safe}.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if fmt == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            _auto_pip("fpdf2")
            try:
                from fpdf import FPDF
            except ImportError:
                return jsonify({"success": False, "msg": "No se pudo instalar fpdf2."}), 500
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=18)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 9, titulo_legible.encode("latin-1", "replace").decode("latin-1"))
        # fpdf2 deja el cursor en el margen derecho tras cada multi_cell:
        # hay que volver al margen izquierdo o el siguiente bloque queda sin ancho.
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "I", 10)
        pdf.multi_cell(0, 7, f"Fecha: {fecha}")
        pdf.set_x(pdf.l_margin)
        pdf.ln(4)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, texto.encode("latin-1", "replace").decode("latin-1"))
        out = pdf.output()
        buf = io.BytesIO(bytes(out))
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=f"{safe}.pdf")

    return jsonify({"success": False, "msg": f"Formato no soportado: {fmt}"}), 400


@app.route("/api/search")
def buscar_transcripciones():
    """Busca un texto dentro de TODAS las transcripciones guardadas."""
    q = (request.args.get("q") or "").strip().lower()
    if not q:
        return jsonify([])
    resultados = []
    try:
        for f in os.listdir(AUDIO_DIR):
            if not f.lower().endswith(".json"):
                continue
            try:
                with open(os.path.join(AUDIO_DIR, f), "r", encoding="utf-8") as fh:
                    data = json.load(fh)
            except Exception:
                continue
            texto = ((data.get("text_es") or "") + " " + (data.get("text_en") or "")).lower()
            title = os.path.splitext(f)[0]
            idx = texto.find(q)
            if q in title.lower() or idx >= 0:
                # Fragmento alrededor de la coincidencia
                snippet = ""
                if idx >= 0:
                    ini = max(0, idx - 60)
                    snippet = ("…" if ini > 0 else "") + texto[ini:idx + len(q) + 80].strip() + "…"
                resultados.append({"title": title, "snippet": snippet})
    except Exception:
        pass
    return jsonify(resultados)


@app.route("/api/upload_url", methods=["POST"])
def upload_desde_url():
    """Descarga el audio de un link (YouTube, etc.) con yt-dlp y lo transcribe."""
    global CURRENT_PROCESS
    if STATE.get("is_processing"):
        return jsonify({"success": False, "msg": "Ya hay una transcripción en proceso."})
    data = request.json or {}
    url = (data.get("url") or "").strip()
    language = data.get("language", "es")
    model_size = (data.get("model") or DEFAULT_MODEL).strip().lower()
    if model_size not in ALLOWED_MODELS:
        model_size = DEFAULT_MODEL
    if not url.lower().startswith(("http://", "https://")):
        return jsonify({"success": False, "msg": "Pega un link válido (http…)."}), 400

    try:
        import yt_dlp
    except ImportError:
        STATE["status_text"] = "Instalando descargador de links (yt-dlp)…"
        _auto_pip("yt-dlp")
        try:
            import yt_dlp
        except ImportError:
            return jsonify({"success": False, "msg": "No se pudo instalar yt-dlp."}), 500

    def descargar_y_transcribir():
        global CURRENT_PROCESS
        try:
            STATE["is_processing"] = True
            STATE["completed"] = False
            STATE["error"] = False
            STATE["progress_percent"] = 2.0
            STATE["status_text"] = "Descargando audio del link…"
            opts = {
                "format": "bestaudio/best",
                "outtmpl": os.path.join(AUDIO_DIR, "%(title).80s.%(ext)s"),
                "quiet": True,
                "noplaylist": True,
                "ffmpeg_location": os.path.dirname(find_binary("ffmpeg")),
            }
            with yt_dlp.YoutubeDL(opts) as ydl:
                info = ydl.extract_info(url, download=True)
                ruta = ydl.prepare_filename(info)
            filename = os.path.basename(ruta)
            STATE["status_text"] = "Audio descargado. Transcribiendo…"
            # is_processing la gestiona proceso_fondo de aquí en adelante
            p = mp.Process(target=proceso_fondo,
                           args=(ruta, filename, STATE, None, None, language, model_size, "transcribe"))
            p.start()
            CURRENT_PROCESS = p
        except Exception as e:
            STATE["is_processing"] = False
            STATE["error"] = True
            STATE["status_text"] = f"❌ No se pudo descargar el link: {e}"

    threading.Thread(target=descargar_y_transcribir, daemon=True).start()
    return jsonify({"success": True})


# ── IA local con Ollama (opcional: solo si está instalado) ──
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://127.0.0.1:11434")
# Preferencia de modelos (si hay varios instalados se usa el primero que exista)
OLLAMA_PREFERIDOS = ["llama3.2", "llama3.1", "llama3", "qwen2.5", "gemma2", "mistral", "phi3"]

OLLAMA_PROMPTS = {
    "resumen": (
        "Eres un asistente editorial en español (colombiano neutro, sin voseo). Resume la siguiente "
        "transcripción: 1-2 oraciones con la idea principal, luego los puntos clave en líneas que "
        "empiezan con '- ', y cierra con la conclusión si la hay. Solo el resumen, sin introducciones."
    ),
    "titulo": (
        "Genera UN título corto y atractivo en español (máximo 8 palabras) para esta transcripción. "
        "Responde SOLO el título, sin comillas ni punto final."
    ),
    "puntos": (
        "Extrae las tareas, acuerdos y acciones mencionadas en esta transcripción como lista de "
        "líneas que empiezan con '- '. Si no hay ninguna, responde 'No se mencionan tareas concretas.'"
    ),
    "post": (
        "Eres un community manager en español. Convierte esta transcripción en un post atractivo para "
        "redes sociales (Instagram/Facebook): gancho inicial, 2-4 párrafos cortos, cierre con llamado "
        "a la acción y 3-5 hashtags. Sin markdown."
    ),
    "devocional": (
        "Eres un editor de contenido cristiano en español. Convierte esta transcripción de sermón en un "
        "devocional breve: título, pasaje bíblico central, reflexión de 3-4 párrafos y una oración final. "
        "Tono cálido y pastoral. Sin markdown."
    ),
    "blog": (
        "Eres un redactor profesional en español. Convierte esta transcripción en un artículo de blog "
        "bien estructurado: título, introducción, desarrollo en secciones y conclusión. Sin markdown, "
        "solo texto con párrafos."
    ),
}


def _ollama_modelos():
    try:
        r = requests.get(f"{OLLAMA_URL}/api/tags", timeout=2)
        if r.status_code == 200:
            return [m.get("name", "") for m in r.json().get("models", [])]
    except Exception:
        pass
    return None  # Ollama no está corriendo


@app.route("/api/ollama/status")
def ollama_status():
    modelos = _ollama_modelos()
    if modelos is None:
        return jsonify({"available": False, "models": []})
    return jsonify({"available": True, "models": modelos})


@app.route("/api/ollama/run", methods=["POST"])
def ollama_run():
    """Ejecuta una acción de IA local sobre una transcripción (o un chat libre)."""
    data = request.json or {}
    accion = data.get("action", "resumen")
    title = data.get("title", "")
    pregunta = (data.get("question") or "").strip()

    modelos = _ollama_modelos()
    if not modelos:
        return jsonify({"success": False, "msg": "Ollama no está corriendo. Instálalo desde ollama.com y descarga un modelo (ej: 'ollama pull llama3.2')."}), 503

    # Elegir modelo: el preferido que esté instalado, si no el primero
    modelo = data.get("ollama_model") or ""
    if not modelo:
        for pref in OLLAMA_PREFERIDOS:
            match = next((m for m in modelos if m.startswith(pref)), None)
            if match:
                modelo = match
                break
        if not modelo:
            modelo = modelos[0]

    jdata, _safe = _leer_json_audio(title)
    texto = ((jdata or {}).get("text_es") or (jdata or {}).get("text_en") or "").strip()
    if not texto:
        return jsonify({"success": False, "msg": "No encontré el texto de ese audio."}), 404

    if accion == "chat":
        if not pregunta:
            return jsonify({"success": False, "msg": "Escribe una pregunta."}), 400
        prompt = (
            "Responde en español (colombiano neutro, sin voseo) usando SOLO la información de esta "
            f"transcripción. Si la respuesta no está en el texto, dilo.\n\nTRANSCRIPCIÓN:\n{texto}\n\n"
            f"PREGUNTA: {pregunta}"
        )
    else:
        base = OLLAMA_PROMPTS.get(accion, OLLAMA_PROMPTS["resumen"])
        prompt = f"{base}\n\nTRANSCRIPCIÓN:\n{texto}"

    try:
        r = requests.post(f"{OLLAMA_URL}/api/generate",
                          json={"model": modelo, "prompt": prompt, "stream": False},
                          timeout=600)
        if r.status_code != 200:
            return jsonify({"success": False, "msg": f"Ollama respondió HTTP {r.status_code}."}), 500
        respuesta = (r.json().get("response") or "").strip()
        return jsonify({"success": True, "text": respuesta, "model": modelo, "action": accion})
    except requests.exceptions.Timeout:
        return jsonify({"success": False, "msg": "La IA local tardó demasiado. Prueba con un modelo más pequeño."}), 504
    except Exception as e:
        return jsonify({"success": False, "msg": str(e)}), 500


PACKAGE_README = """TRANSCRIPTOR DE AUDIOS — Instalación en Mac
============================================

1. Descomprime esta carpeta donde quieras (por ejemplo, el Escritorio).
2. Haz doble clic en "Transcriptor_Facil.command".
   - La primera vez instala Python, FFmpeg y la IA (Whisper). Puede tardar
     varios minutos y descargar el modelo "medium" (~1.5 GB). Ten paciencia.
   - Si macOS bloquea el archivo: clic derecho -> Abrir, o en Terminal:
       chmod +x Transcriptor_Facil.command
3. Se abrira tu navegador en http://127.0.0.1:5111

MODELO DE TRANSCRIPCION
-----------------------
Por defecto usa "medium" (muy buena calidad, ~5 GB de RAM). Puedes cambiarlo:
- Desde la interfaz, en el selector "Modelo (calidad)" al subir un audio.
- O de forma permanente: abre "Transcriptor_Facil.command" con un editor de
  texto y cambia el valor de WHISPER_MODEL por: tiny / base / small / medium / large-v3

ACTUALIZAR
----------
No necesitas volver a descargar el paquete: abre el engranaje (Configuracion) y
pulsa "Buscar actualizacion". Descarga la ultima version desde GitHub. Si avisa
que reinicies, cierra la ventana negra (Terminal) y reabre "Transcriptor_Facil.command".

NOTAS
-----
- Es 100% local: tus audios no salen de tu Mac.
- Para apagarlo, cierra la ventana negra (Terminal).
"""


APP_README = """TRANSCRIPTOR — App para Mac (sin Terminal)
============================================

1. Descomprime este .zip → te queda "Transcriptor.app".
2. (Recomendado) Arrastra "Transcriptor.app" a tu carpeta Aplicaciones.
3. Haz doble clic en Transcriptor.
   - Si macOS lo bloquea ("desarrollador no identificado"): clic derecho ->
     Abrir -> Abrir. Solo la primera vez.
   - La primera vez instala la IA por dentro (sin ventana negra). Puede tardar
     varios minutos y descargar el modelo (~1.5 GB). Espera la notificacion "Listo".
4. Se abre en su propia ventana. Listo.

- Es 100% local: tus audios no salen del Mac.
- Modelo por defecto: medium. Lo cambias en el selector "Modelo (calidad)".
- Para actualizar: dentro de la app, engranaje -> "Buscar actualizacion".
"""


@app.route("/api/package/app")
def descargar_app():
    """Genera al vuelo un .zip con Transcriptor.app (app nativa, sin Terminal),
    configurada con modelo medium. Siempre refleja el codigo actual."""
    import io, zipfile
    ROOT = os.path.dirname(os.path.abspath(__file__))
    APP = "Transcriptor.app"
    RES = f"{APP}/Contents/Resources/app"
    tpl = os.path.join(ROOT, "tools", "app_template")
    icns = os.path.join(ROOT, "assets", "icon.icns")

    def leer(p):
        with open(p, "r", encoding="utf-8") as f:
            return f.read()

    def leer_b(p):
        with open(p, "rb") as f:
            return f.read()

    if not os.path.exists(os.path.join(tpl, "launcher.sh")) or not os.path.exists(icns):
        return jsonify({"success": False, "msg": "Faltan plantillas de la app en este servidor."}), 500

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # Estructura del bundle
        z.writestr(f"{APP}/Contents/Info.plist", leer(os.path.join(tpl, "Info.plist")))
        zl = zipfile.ZipInfo(f"{APP}/Contents/MacOS/Transcriptor"); zl.external_attr = 0o755 << 16
        z.writestr(zl, leer(os.path.join(tpl, "launcher.sh")))
        z.writestr(f"{APP}/Contents/Resources/icon.icns", leer_b(icns))

        # Codigo de la app
        z.writestr(f"{RES}/server.py", leer(os.path.join(ROOT, "server.py")))
        zr = zipfile.ZipInfo(f"{RES}/run_server.sh"); zr.external_attr = 0o755 << 16
        z.writestr(zr, leer(os.path.join(ROOT, "run_server.sh")))
        req = os.path.join(ROOT, "requirements.txt")
        if os.path.exists(req):
            z.writestr(f"{RES}/requirements.txt", leer(req))

        fe_root = os.path.join(ROOT, "frontend", "public")
        for dirpath, dirnames, filenames in os.walk(fe_root):
            dirnames[:] = [d for d in dirnames if d != "downloads" and not d.startswith(".")]
            for fn in filenames:
                if fn.startswith(".") or fn.lower().endswith(".zip"):
                    continue
                full = os.path.join(dirpath, fn)
                rel = os.path.relpath(full, fe_root)
                z.write(full, f"{RES}/frontend/public/{rel}")

        # Plantillas + icono dentro (para poder regenerar la app desde dentro)
        z.writestr(f"{RES}/tools/app_template/Info.plist", leer(os.path.join(tpl, "Info.plist")))
        z.writestr(f"{RES}/tools/app_template/launcher.sh", leer(os.path.join(tpl, "launcher.sh")))
        z.writestr(f"{RES}/assets/icon.icns", leer_b(icns))

        z.writestr("LEEME_PRIMERO.txt", APP_README)

    buf.seek(0)
    return send_file(buf, mimetype="application/zip", as_attachment=True,
                     download_name="Transcriptor-App-Mac.zip")


@app.route("/api/package/mac")
def descargar_paquete_mac():
    """Genera al vuelo un .zip con la app lista para correr en otro Mac,
    con el modelo por defecto en 'medium'. Siempre refleja el codigo actual."""
    import io, zipfile
    ROOT = os.path.dirname(os.path.abspath(__file__))
    TOP = "TranscriptorAudios"

    def leer(path):
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # server.py
        z.writestr(f"{TOP}/server.py", leer(os.path.join(ROOT, "server.py")))

        # run_server.sh con WHISPER_MODEL=medium por defecto (override por env)
        run_sh = leer(os.path.join(ROOT, "run_server.sh"))
        run_sh = run_sh.replace(
            "python3 server.py",
            'export WHISPER_MODEL="${WHISPER_MODEL:-medium}"\npython3 server.py'
        )
        zi = zipfile.ZipInfo(f"{TOP}/run_server.sh"); zi.external_attr = 0o755 << 16
        z.writestr(zi, run_sh)

        # .command que abre localhost (no la URL de la nube)
        cmd = leer(os.path.join(ROOT, "Transcriptor_Facil.command"))
        cmd = cmd.replace("https://transcriptor-audios.web.app", "http://127.0.0.1:5111")
        zic = zipfile.ZipInfo(f"{TOP}/Transcriptor_Facil.command"); zic.external_attr = 0o755 << 16
        z.writestr(zic, cmd)

        # requirements.txt
        req_path = os.path.join(ROOT, "requirements.txt")
        if os.path.exists(req_path):
            z.writestr(f"{TOP}/requirements.txt", leer(req_path))

        # frontend/public (sin la carpeta downloads ni ocultos)
        fe_root = os.path.join(ROOT, "frontend", "public")
        for dirpath, dirnames, filenames in os.walk(fe_root):
            dirnames[:] = [d for d in dirnames if d != "downloads" and not d.startswith(".")]
            for fn in filenames:
                if fn.startswith(".") or fn.lower().endswith(".zip"):
                    continue
                full = os.path.join(dirpath, fn)
                rel = os.path.relpath(full, fe_root)
                z.write(full, f"{TOP}/frontend/public/{rel}")

        # Instrucciones
        z.writestr(f"{TOP}/LEEME_PRIMERO.txt", PACKAGE_README)

    buf.seek(0)
    return send_file(buf, mimetype="application/zip", as_attachment=True,
                     download_name="TranscriptorAudios-Mac.zip")


@app.route("/api/update", methods=["POST"])
def api_update():
    """Descarga los archivos más recientes del repo en GitHub y reemplaza los
    locales. El frontend aplica al recargar; si cambia el server hay que reiniciar."""
    ROOT = os.path.dirname(os.path.abspath(__file__))
    updated, errors = [], []
    server_changed = False
    headers = {"Accept": "application/vnd.github.raw", "User-Agent": "transcriptor-updater"}
    for rel in UPDATE_FILES:
        url = f"{GITHUB_API}/{rel}?ref=main"
        try:
            r = requests.get(url, headers=headers, timeout=25)
            if r.status_code != 200:
                errors.append(f"{rel}: HTTP {r.status_code}")
                continue
            new = r.content
            dest = os.path.join(ROOT, *rel.split("/"))
            old = b""
            if os.path.exists(dest):
                with open(dest, "rb") as f:
                    old = f.read()
            if new != old:
                os.makedirs(os.path.dirname(dest), exist_ok=True)
                # Escritura segura: archivo temporal + reemplazo atómico
                tmp = dest + ".new"
                with open(tmp, "wb") as f:
                    f.write(new)
                os.replace(tmp, dest)
                updated.append(rel)
                if rel in ("server.py", "run_server.sh", "requirements.txt"):
                    server_changed = True
        except Exception as e:
            errors.append(f"{rel}: {e}")
    return jsonify({
        "success": len(errors) == 0,
        "updated": updated,
        "errors": errors,
        "needs_restart": server_changed,
    })


if __name__ == "__main__":
    try:
        if mp.get_start_method(allow_none=True) != 'spawn':
            mp.set_start_method('spawn', force=True)
    except RuntimeError: pass

    manager = mp.Manager()
    STATE = manager.dict({
        "is_processing": False,
        "status_text": "Listo para transcribir.",
        "is_installing": False,
        "progress_percent": 0.0,
        "completed": False,
        "error": False,
        "current_filename": "",
        "title": ""
    })

    limpiar_temporales()

    # Vigilante: apaga el server cuando se cierra la ventana (libera memoria)
    if AUTO_SHUTDOWN:
        threading.Thread(target=_watchdog, daemon=True).start()
        print("🛡️  Apagado automático al cerrar la ventana: ACTIVADO")

    print("🚀 Servidor listo en http://127.0.0.1:5111")
    # Cambiado a 0.0.0.0 para acceso desde otros dispositivos y estabilidad
    app.run(host="0.0.0.0", port=5111, debug=False, threaded=True)
